import docx

from .results import ContainerResult, FigureResult, TableResult

class WordGenerator(object):
    def __init__(self, results, results_dir):
        self.results_dir = results_dir
        self.results = results
        self.doc = docx.Document()
        self._written_ids = set()
        
    def generate(self, output_file, result_ids=None):
        if result_ids is None:
            result_ids = self.results['root'].children
        for result_id in result_ids:
            self._generate(result_id, 0)
        self.doc.save(output_file)
            
    def _generate(self, result_id, heading_level):
        if result_id in self._written_ids:
            return
        self._written_ids.add(result_id)
        result = self.results[result_id]
        self.doc.add_heading(result.name, heading_level)
        if isinstance(result, ContainerResult):
            for child_id in result.children:
                self._generate(child_id, heading_level+1)
        elif isinstance(result, FigureResult):
            self.doc.add_picture(result.full_path, width=docx.shared.Inches(6.0))
        elif isinstance(result, TableResult):
            if result.pre != '':
                self.doc.add_paragraph(result.pre)
            table = self.doc.add_table(1, len(result.headings)) # headings also count as rows
            table.style.font.bold = True
            header_cells = table.rows[0].cells
            for header_cell, heading in zip(header_cells, result.headings):
                header_cell.text = heading
            for row in result.rows:
                row_cells = table.add_row().cells
                for row_cell, cell_text in zip(row_cells, row):
                    row_cell.text = cell_text
            if result.post != '':
                self.doc.add_paragraph(result.post)
        else:
            raise NotImplementedError('Result type: ' + result.result_type)
